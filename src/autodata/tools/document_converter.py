"""Document converter for exam files.

Supports DOCX, DOC, PDF, and image files with provenance tracking.
"""

from __future__ import annotations

import hashlib
import json
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional


def extract_docx(file_path: Path) -> list[dict]:
    """Extract text from DOCX file using python-docx.

    Returns list of text blocks with provenance.
    """
    blocks = []
    try:
        from docx import Document
        doc = Document(str(file_path))

        para_idx = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            blocks.append({
                "block_id": f"tb_{hashlib.md5(f'{file_path.name}:p0:para{i}'.encode()).hexdigest()[:12]}",
                "source_file": file_path.name,
                "page_number": 0,
                "paragraph_id": para_idx,
                "table_id": "",
                "text": text,
                "extraction_method": "python-docx",
                "content_hash": hashlib.md5(text.encode()).hexdigest(),
            })
            para_idx += 1

        # Extract tables
        for t_idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            text = "\n".join(table_text)
            if text.strip():
                blocks.append({
                    "block_id": f"tb_{hashlib.md5(f'{file_path.name}:table{t_idx}'.encode()).hexdigest()[:12]}",
                    "source_file": file_path.name,
                    "page_number": 0,
                    "paragraph_id": 0,
                    "table_id": f"table_{t_idx}",
                    "text": text,
                    "extraction_method": "python-docx-table",
                    "content_hash": hashlib.md5(text.encode()).hexdigest(),
                })

    except ImportError:
        blocks.append({
            "block_id": "error_docx_import",
            "source_file": file_path.name,
            "page_number": 0,
            "paragraph_id": 0,
            "table_id": "",
            "text": f"ERROR: python-docx not installed",
            "extraction_method": "error",
            "content_hash": "",
        })
    except Exception as e:
        blocks.append({
            "block_id": "error_docx_parse",
            "source_file": file_path.name,
            "page_number": 0,
            "paragraph_id": 0,
            "table_id": "",
            "text": f"ERROR: {str(e)[:200]}",
            "extraction_method": "error",
            "content_hash": "",
        })

    return blocks


def extract_doc(file_path: Path) -> list[dict]:
    """Extract text from DOC file.

    Tries: 1) python-docx (for mislabeled OOXML), 2) LibreOffice, 3) antiword.
    """
    # First try python-docx (handles mislabeled OOXML files)
    try:
        from docx import Document
        doc = Document(str(file_path))
        if doc.paragraphs:
            blocks = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                blocks.append({
                    "block_id": f"tb_{hashlib.md5(f'{file_path.name}:p0:para{i}'.encode()).hexdigest()[:12]}",
                    "source_file": file_path.name,
                    "page_number": 0,
                    "paragraph_id": i,
                    "table_id": "",
                    "text": text,
                    "extraction_method": "python-docx-for-doc",
                    "content_hash": hashlib.md5(text.encode()).hexdigest(),
                })
            if blocks:
                return blocks
    except Exception:
        pass

    # Try LibreOffice headless conversion
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, str(file_path)],
                capture_output=True, text=True, timeout=60
            )
            txt_files = list(Path(tmpdir).glob("*.txt"))
            if txt_files:
                text = txt_files[0].read_text(encoding="utf-8", errors="replace")
                return _split_text_to_blocks(text, file_path.name, "libreoffice")
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Try antiword
    try:
        result = subprocess.run(
            ["antiword", str(file_path)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return _split_text_to_blocks(result.stdout, file_path.name, "antiword")
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Try catdoc
    try:
        result = subprocess.run(
            ["catdoc", str(file_path)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return _split_text_to_blocks(result.stdout, file_path.name, "catdoc")
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    return [{
        "block_id": "error_doc_extraction",
        "source_file": file_path.name,
        "page_number": 0,
        "paragraph_id": 0,
        "table_id": "",
        "text": f"ERROR: Could not extract DOC file with any method",
        "extraction_method": "error",
        "content_hash": "",
    }]


def extract_pdf(file_path: Path) -> list[dict]:
    """Extract text from PDF file.

    Tries text extraction first, falls back to OCR if needed.
    """
    # Try pdftotext first
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", str(file_path), "-"],
            capture_output=True, text=True, timeout=60
        )
        text = result.stdout.strip()
        if len(text) > 50:  # Meaningful text extracted
            return _split_text_to_blocks(text, file_path.name, "pdftotext")
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Check if scanned (try first page)
    try:
        result = subprocess.run(
            ["pdftotext", str(file_path), "-", "-l", "1"],
            capture_output=True, text=True, timeout=10
        )
        if len(result.stdout.strip()) < 20:
            # Likely scanned, would need OCR
            return [{
                "block_id": "scanned_pdf_needs_ocr",
                "source_file": file_path.name,
                "page_number": 1,
                "paragraph_id": 0,
                "table_id": "",
                "text": "SCANNED_PDF: Requires OCR processing",
                "extraction_method": "scanned_detection",
                "content_hash": "",
            }]
    except Exception:
        pass

    return [{
        "block_id": "error_pdf_extraction",
        "source_file": file_path.name,
        "page_number": 0,
        "paragraph_id": 0,
        "table_id": "",
        "text": f"ERROR: Could not extract PDF text",
        "extraction_method": "error",
        "content_hash": "",
    }]


def extract_image_ocr(file_path: Path) -> list[dict]:
    """Extract text from image using OCR."""
    try:
        result = subprocess.run(
            ["tesseract", str(file_path), "stdout", "-l", "chi_sim+eng"],
            capture_output=True, text=True, timeout=60
        )
        if result.stdout.strip():
            return _split_text_to_blocks(result.stdout, file_path.name, "tesseract-ocr")
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    return [{
        "block_id": "error_ocr",
        "source_file": file_path.name,
        "page_number": 0,
        "paragraph_id": 0,
        "table_id": "",
        "text": f"ERROR: OCR failed or tesseract not installed",
        "extraction_method": "error",
        "content_hash": "",
    }]


def _split_text_to_blocks(text: str, source_file: str, method: str) -> list[dict]:
    """Split text into paragraph blocks with provenance."""
    blocks = []
    paragraphs = text.split("\n\n")

    for i, para in enumerate(paragraphs):
        para = para.strip()
        if not para:
            continue
        blocks.append({
            "block_id": f"tb_{hashlib.md5(f'{source_file}:p0:para{i}'.encode()).hexdigest()[:12]}",
            "source_file": source_file,
            "page_number": 0,
            "paragraph_id": i,
            "table_id": "",
            "text": para,
            "extraction_method": method,
            "content_hash": hashlib.md5(para.encode()).hexdigest(),
        })

    return blocks


def convert_document(file_path: Path) -> list[dict]:
    """Convert a document to text blocks.

    Dispatches to appropriate extractor based on file type.
    """
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return extract_docx(file_path)
    elif suffix == ".doc":
        return extract_doc(file_path)
    elif suffix == ".pdf":
        return extract_pdf(file_path)
    elif suffix in (".jpg", ".jpeg", ".png", ".tiff", ".bmp"):
        return extract_image_ocr(file_path)
    elif suffix == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return _split_text_to_blocks(text, file_path.name, "text-read")
    else:
        return [{
            "block_id": "error_unsupported",
            "source_file": file_path.name,
            "page_number": 0,
            "paragraph_id": 0,
            "table_id": "",
            "text": f"ERROR: Unsupported file type: {suffix}",
            "extraction_method": "error",
            "content_hash": "",
        }]

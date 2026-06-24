"""Exam data inventory for Phase 4.

Scans exam_raw_data/ and classifies each file.
"""

from __future__ import annotations

import json
import os
import subprocess
import time
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[2].parent


def detect_file_type(file_path: Path) -> str:
    """Detect actual file type using file command."""
    try:
        result = subprocess.run(
            ["file", "--mime-type", "-b", str(file_path)],
            capture_output=True, text=True, timeout=10
        )
        mime = result.stdout.strip()
        if "pdf" in mime:
            return "pdf"
        elif "msword" in mime or "officedocument" in mime or "zip" in mime:
            return "docx"
        elif "ole2" in mime or "msword" in mime:
            return "doc"
        elif "image" in mime:
            return "image"
        elif "text" in mime:
            return "text"
        return file_path.suffix.lower().lstrip(".")
    except Exception:
        return file_path.suffix.lower().lstrip(".")


def estimate_page_count(file_path: Path, file_type: str) -> int:
    """Estimate page count without full parsing."""
    size_kb = file_path.stat().st_size / 1024
    if file_type == "pdf":
        try:
            result = subprocess.run(
                ["pdfinfo", str(file_path)],
                capture_output=True, text=True, timeout=10
            )
            for line in result.stdout.splitlines():
                if line.startswith("Pages:"):
                    return int(line.split(":")[1].strip())
        except Exception:
            pass
        return max(1, int(size_kb / 100))
    elif file_type in ("docx", "doc"):
        return max(1, int(size_kb / 20))
    return 1


def check_answer_key(file_path: Path) -> bool:
    """Check if filename suggests answer key is included."""
    name = file_path.stem.lower()
    return any(kw in name for kw in ["答案", "answer", "含答案", "带答案", "试题及"])


def check_scanned(file_path: Path, file_type: str) -> bool:
    """Check if file is scanned/image-based."""
    if file_type == "pdf":
        try:
            result = subprocess.run(
                ["pdftotext", str(file_path), "-", "-l", "1"],
                capture_output=True, text=True, timeout=10
            )
            text = result.stdout.strip()
            if len(text) < 20:
                return True
        except Exception:
            pass
    return file_type == "image"


def check_language(file_path: Path) -> str:
    """Detect primary language from filename."""
    name = file_path.name
    chinese_chars = sum(1 for c in name if '\u4e00' <= c <= '\u9fff')
    if chinese_chars > 2:
        return "zh"
    return "en"


def check_formulas_tables(file_path: Path, file_type: str) -> dict:
    """Quick check for formulas/tables presence."""
    result = {"formulas": False, "tables": False, "images": False}
    if file_type in ("docx", "doc"):
        try:
            result2 = subprocess.run(
                ["python3", "-c", f"""
import sys
try:
    from docx import Document
    doc = Document('{file_path}')
    text = ' '.join(p.text for p in doc.paragraphs)
    import re
    formulas = bool(re.search(r'[=∑∫√αβγδ]|\$.*?\$|\\\\[a-zA-Z]+', text))
    tables = len(doc.tables) > 0
    print(f'formulas={{formulas}},tables={{tables}}')
except Exception as e:
                    print(f'error={{e}}')
"""],
                capture_output=True, text=True, timeout=15
            )
            if "formulas=True" in result2.stdout:
                result["formulas"] = True
            if "tables=True" in result2.stdout:
                result["tables"] = True
        except Exception:
            pass
    return result


def get_parsing_method(file_type: str, scanned: bool) -> str:
    """Determine parsing method."""
    if scanned:
        return "ocr"
    if file_type == "docx":
        return "python-docx"
    if file_type == "doc":
        return "libreoffice_or_antiword"
    if file_type == "pdf":
        return "pdftotext_or_ocr"
    if file_type == "image":
        return "ocr"
    return "unknown"


def get_risk_level(file_type: str, scanned: bool, size_kb: float) -> str:
    """Assess parsing risk level."""
    if scanned:
        return "high"
    if file_type == "doc" and size_kb > 500:
        return "medium"
    if file_type == "pdf":
        return "medium"
    return "low"


def build_inventory(exam_dir: Path) -> list[dict]:
    """Build inventory of all exam files."""
    inventory = []

    for file_path in sorted(exam_dir.iterdir()):
        if file_path.is_dir() or file_path.name.startswith("."):
            continue

        file_type = detect_file_type(file_path)
        size_kb = file_path.stat().st_size / 1024
        pages = estimate_page_count(file_path, file_type)
        scanned = check_scanned(file_path, file_type)
        language = check_language(file_path)
        has_answer_key = check_answer_key(file_path)
        content_check = check_formulas_tables(file_path, file_type)
        parsing_method = get_parsing_method(file_type, scanned)
        risk_level = get_risk_level(file_type, scanned, size_kb)

        entry = {
            "file_name": file_path.name,
            "file_path": str(file_path),
            "file_type": file_type,
            "file_size_kb": round(size_kb, 1),
            "estimated_pages": pages,
            "scanned": scanned,
            "language": language,
            "has_answer_key": has_answer_key,
            "contains_formulas": content_check.get("formulas", False),
            "contains_tables": content_check.get("tables", False),
            "contains_images": content_check.get("images", False),
            "parsing_method": parsing_method,
            "risk_level": risk_level,
        }
        inventory.append(entry)

    return inventory


def save_inventory(inventory: list[dict]) -> tuple[Path, Path]:
    """Save inventory as JSON and MD."""
    report_dir = PROJECT_ROOT / "data" / "reports" / "phase_4_exam_extraction"
    report_dir.mkdir(parents=True, exist_ok=True)

    json_path = report_dir / "exam_inventory.json"
    md_path = report_dir / "exam_inventory.md"

    # Save JSON
    with open(json_path, "w") as f:
        json.dump(inventory, f, indent=2, ensure_ascii=False)

    # Save MD
    with open(md_path, "w") as f:
        f.write("# 考试数据清单\n\n")
        f.write(f"共 {len(inventory)} 个文件\n\n")

        # Summary
        types = {}
        for item in inventory:
            t = item["file_type"]
            types[t] = types.get(t, 0) + 1
        f.write("## 文件类型统计\n\n")
        for t, count in sorted(types.items()):
            f.write(f"- {t}: {count}\n")

        scanned_count = sum(1 for item in inventory if item["scanned"])
        answer_count = sum(1 for item in inventory if item["has_answer_key"])
        f.write(f"\n- 扫描文件: {scanned_count}\n")
        f.write(f"- 含答案文件: {answer_count}\n")

        f.write("\n## 文件详情\n\n")
        f.write("| 文件名 | 类型 | 大小(KB) | 页数 | 扫描 | 语言 | 答案 | 解析方法 | 风险 |\n")
        f.write("|--------|------|----------|------|------|------|------|----------|------|\n")
        for item in inventory:
            scanned = "是" if item["scanned"] else "否"
            answer = "是" if item["has_answer_key"] else "否"
            f.write(f"| {item['file_name']} | {item['file_type']} | {item['file_size_kb']} | "
                    f"{item['estimated_pages']} | {scanned} | {item['language']} | {answer} | "
                    f"{item['parsing_method']} | {item['risk_level']} |\n")

    return json_path, md_path
